import os
import json
import io
import re
import PyPDF2
import docx
import tempfile
import fitz  # PyMuPDF for PDF handling
import docx  # python-docx for Word docs
import uvicorn
import random
import shutil
import google.generativeai as genai
from dotenv import load_dotenv
from pymongo import MongoClient
from bson.objectid import ObjectId
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from bson import ObjectId
from PIL import Image
from dotenv import load_dotenv
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import bson.errors

# Load environment variables
load_dotenv()

# Configure Google Gemini API
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY environment variable is not set")

genai.configure(api_key=GOOGLE_API_KEY)

# MongoDB Connection Configuration
mongo_client = MongoClient("mongodb://localhost:27017/")
db = mongo_client["employeeDB"]
employees_collection = db["employees"]

# Configure model
generation_config = {
    "temperature": 0.2,
    "top_p": 0.8,
    "top_k": 40,
    "max_output_tokens": 1024,
}

model = genai.GenerativeModel(
    model_name="gemini-2.0-flash",
    generation_config=generation_config
)

# Create FastAPI app
app = FastAPI(title="Job Description Skills Matcher API")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Update this in production to specific origins
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Define response models
class Employee(BaseModel):
    id: str
    name: str
    role: str
    location: str
    image_url: str
    skills: List[str]
    match: float = 0
    match_percentage: str = "0%"
    match_justification: str = ""
    experience: int = 0
    salaryRange: str = ""

class MatchResponse(BaseModel):
    job_description_skills: dict
    candidates_count: int
    candidates_updated: int
    matched_candidates: List[dict]

def generate_match_score():
    """Generate a random match score between 70 and 100."""
    return random.randint(70, 100)

class DocumentProcessor:
    def __init__(self):
        # Get API key from environment variable
        self.api_key = os.environ.get("GOOGLE_API_KEY")
        if not self.api_key:
            raise EnvironmentError("GOOGLE_API_KEY environment variable is not set")
        
        # Configure Google Gemini API
        genai.configure(api_key=self.api_key)
        
        # Initialize the model - Gemini 1.5 Flash
        self.model = genai.GenerativeModel('gemini-2.0-flash')
        
        # MongoDB collection is passed from the outer scope
        self.collection = employees_collection
    
    def process_document(self, file_path):
        """Process a document file and extract personal information"""
        # Determine file type
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._process_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            return self._process_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _process_pdf(self, file_path):
        """Process PDF file using Gemini 1.5"""
        # Open the PDF
        pdf_document = fitz.open(file_path)
        
        # Extract text from all pages
        full_text = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            full_text += page.get_text()
        
        # Create an image of the first page
        first_page = pdf_document[0]
        pix = first_page.get_pixmap()
        img_bytes = pix.tobytes("png")
        img = Image.open(io.BytesIO(img_bytes))
        
        # Use Gemini to extract information from both text and image
        return self._extract_with_gemini(full_text, img)
    
    def _process_word(self, file_path):
        """Process Word document"""
        if file_path.endswith('.docx'):
            # Extract text from Word document
            doc = docx.Document(file_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Use Gemini to extract information
            return self._extract_with_gemini(full_text)
        else:
            # For .doc files
            raise NotImplementedError("DOC file processing not implemented")
    
    def _extract_with_gemini(self, text, image=None):
        """Extract personal information using Google Gemini 1.5"""
        # Create prompt for Gemini
        prompt = """
        Extract the following details from the document if present:
        1. Full name
        2. Role or job title
        3. Email address
        4. Phone number
        5. Years of experience
        6. Location (City, State)
        7. List of skills
        8. Education qualification
        9. Languages spoken with proficiency level

        Format the response as a JSON object with these fields:
        {
            "name": "",
            "role": "",
            "email": "",
            "phone": "",
            "experience": 0,
            "location": "",
            "skills": [],
            "education": "",
            "languages": []
        }

        If a field is missing, return an empty string or appropriate empty value.
        """
        
        try:
            # If we have both text and image
            if image:
                response = self.model.generate_content([prompt, text, image])
            else:
                response = self.model.generate_content([prompt, text])
            
            # Get the response text
            response_text = response.text
            
            # Extract JSON from response (handle potential extra text)
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            
            if json_start >= 0 and json_end > 0:
                json_str = response_text[json_start:json_end]
                extracted_info = json.loads(json_str)
            else:
                # Fallback to regex if JSON parsing fails
                extracted_info = self._extract_with_regex(text)
                
            # Add document text 
            return extracted_info
            
        except Exception as e:
            print(f"Error with Gemini extraction: {e}")
            # Fallback to regex-based extraction
            return self._extract_with_regex(text)
    
    def _extract_with_regex(self, text):
        """Fallback method: Extract personal information using regex patterns"""
        # Initialize personal info dictionary
        personal_info = {
            "name": "",
            "role": "",
            "email": "",
            "phone": "",
            "experience": 0,
            "location": "",
            "skills": [],
            "education": "",
            "languages": []
        }
        
        # Extract email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            personal_info["email"] = emails[0]
        
        # Extract phone numbers
        phone_pattern = r'\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'
        phones = re.findall(phone_pattern, text)
        if phones:
            personal_info["phone"] = phones[0]
        
        # Extract potential name
        name_patterns = [
            r'(?:name|full name|I am)[:\s]+([A-Z][a-z]+(?:\s[A-Z][a-z]+){1,3})',
            r'(?:Mr\.|Mrs\.|Ms\.|Dr\.)\s([A-Z][a-z]+(?:\s[A-Z][a-z]+){1,3})'
        ]
        
        for pattern in name_patterns:
            name_match = re.search(pattern, text, re.IGNORECASE)
            if name_match:
                personal_info["name"] = name_match.group(1)
                break
        
        return personal_info
    
    def save_to_mongodb(self, extracted_info):
        image_url = "no-image-svgrepo-com.svg" 
        formatted_data = {
            "name": extracted_info.get("name", ""),
            "role": extracted_info.get("role", ""),
            "email": extracted_info.get("email", ""),
            "phone": extracted_info.get("phone", ""),
            "experience": extracted_info.get("experience", 0),
            "location": extracted_info.get("location", ""),
            "skills": extracted_info.get("skills", []),
            "education": extracted_info.get("education", ""),
            "languages": extracted_info.get("languages", []),
            "match": generate_match_score(),
            "image_url": image_url
        }
        
        result = self.collection.insert_one(formatted_data)
        return str(result.inserted_id)
def extract_text_from_pdf(file_content):
    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(file_content):
    doc = docx.Document(io.BytesIO(file_content))
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def extract_skills_with_gemini(text):
    prompt = """
    Extract technical skills, soft skills, and relevant experience from the following job description.
    Format the output as a JSON with the following format:
    {
        "technical_skills": ["skill1", "skill2", ...],
        "soft_skills": ["skill1", "skill2", ...],
        "years_experience_required": "X years" (if mentioned, otherwise null)
    }
    Only include clearly mentioned skills. Do not infer skills that aren't explicitly stated.
    Here's the job description:
    
    """
    
    response = model.generate_content(prompt + text)
    
    try:
        # Extract the JSON content from the response
        content_text = response.text
        if "```json" in content_text:
            json_str = content_text.split("```json")[1].split("```")[0].strip()
        elif "```" in content_text:
            json_str = content_text.split("```")[1].strip()
        else:
            json_str = content_text.strip()
            
        skills_data = json.loads(json_str)
        return skills_data
    except Exception as e:
        return {"error": f"Failed to parse skills from job description: {str(e)}"}

async def process_job_description(file_content, file_extension):
    try:
        # Extract text based on file type
        if file_extension.lower() == '.pdf':
            text = extract_text_from_pdf(file_content)
        elif file_extension.lower() in ['.docx', '.doc']:
            text = extract_text_from_docx(file_content)
        else:
            return {"error": "Unsupported file format. Please upload a PDF or Word document."}
        
        # Extract skills using Gemini
        skills = extract_skills_with_gemini(text)
        
        return skills
    except Exception as e:
        return {"error": f"Error processing job description: {str(e)}"}

def match_candidate(jd_skills, candidate_data):
    try:
        if "error" in candidate_data:
            return candidate_data
            
        # Extract required skills from the input
        required_tech_skills = jd_skills.get("technical_skills", [])
        required_soft_skills = jd_skills.get("soft_skills", [])
        
        # Get candidate skills
        candidate_skills = candidate_data.get("skills", [])
        
        # Handle empty skills
        if not candidate_skills:
            return {**candidate_data, "error": "No skills found for candidate"}
            
        # Calculate match score
        tech_match_percent = sum(1 for skill in required_tech_skills if any(skill.lower() in c_skill.lower() or c_skill.lower() in skill.lower() for c_skill in candidate_skills)) / len(required_tech_skills) * 100 if required_tech_skills else 0
        soft_match_percent = sum(1 for skill in required_soft_skills if any(skill.lower() in c_skill.lower() or c_skill.lower() in skill.lower() for c_skill in candidate_skills)) / len(required_soft_skills) * 100 if required_soft_skills else 0
        
        if tech_match_percent == 0:
            soft_match_percent = 0
        else:
            score = (tech_match_percent * 0.9) + (soft_match_percent * 0.1) #90% Technical skills, 10% soft skills
        
        # Extract years_experience_required and convert to integer
        years_required = 0
        if "years_experience_required" in jd_skills and jd_skills["years_experience_required"]:
            # Try to extract the number from the string like "3 years" or "5+ years"
            exp_text = jd_skills["years_experience_required"]
            try:
                years_required = int(''.join(filter(str.isdigit, exp_text)))
            except:
                years_required = 0

        # Calculate the experience bonus (10% if candidate meets or exceeds required experience)
        experience_bonus = 0
        if tech_match_percent > 0 and years_required > 0:
            candidate_exp = candidate_data.get("experience", 0)
            if candidate_exp >= years_required:
                experience_bonus = 10
            else:
                # Partial bonus based on percentage of required experience
                experience_bonus = (candidate_exp / years_required) * 100
        experience_bonus = min(10, experience_bonus)

        # Add the experience bonus to the score
        final_score = score + experience_bonus
        final_score = min(100, final_score)  # Cap at 100%
        final_score = round(final_score, 1)
        
        # Create a copy of the candidate data and update the match score
        result = candidate_data.copy()
        result["match"] = final_score
        result["match_percentage"] = f"{final_score:.1f}%"
        result["match_justification"] = f"Matches {tech_match_percent:.1f}% of technical skills and {soft_match_percent:.1f}% of soft skills. Experience bonus: {experience_bonus:.1f}%."
        
        return result
    except Exception as e:
        return {**candidate_data, "error": f"Error matching candidate: {str(e)}"}

def get_candidates_from_mongodb():
    try:
        # Get all candidates from MongoDB
        candidates = list(employees_collection.find({}))
        
        # Verify we got candidates
        count = len(candidates)
        if count == 0:
            print("WARNING: No candidates found in MongoDB database")
            
        # Convert ObjectId to string for JSON serialization
        for candidate in candidates:
            if '_id' in candidate:
                candidate['id'] = str(candidate['_id'])  # Add id field for frontend compatibility
                candidate['_id'] = str(candidate['_id'])
                
        return candidates
    except Exception as e:
        print(f"MongoDB Error: {str(e)}")
        return {"error": f"Error retrieving candidates from MongoDB: {str(e)}"}

def update_candidate_in_mongodb(candidate_id, match_data):
    try:
        # Convert string ID back to ObjectId
        obj_id = ObjectId(candidate_id)
        
        # Update only the match-related fields
        result = employees_collection.update_one(
            {"_id": obj_id},
            {"$set": {
                "match": match_data["match"],
                "match_percentage": match_data["match_percentage"],
                "match_justification": match_data["match_justification"]
            }}
        )
        
        # Check if update was successful
        return result.modified_count > 0
    except Exception as e:
        print(f"Error updating candidate in MongoDB: {str(e)}")
        return False

def get_mongodb_stats():
    try:
        # Gather information about the MongoDB database and collection
        db_names = mongo_client.list_database_names()
        collection_names = db.list_collection_names() if "employeedb" in db_names else []
        employee_count = employees_collection.count_documents({}) if "employees" in collection_names else 0
        
        # Get a sample document (for debugging)
        sample = None
        if employee_count > 0:
            sample = employees_collection.find_one()
            if sample and '_id' in sample:
                sample['_id'] = str(sample['_id'])
        
        return {
            "connection_successful": True,
            "databases": db_names,
            "collections": collection_names,
            "employee_count": employee_count,
            "sample_document": sample
        }
    except Exception as e:
        return {
            "connection_successful": False,
            "error": str(e)
        }

async def batch_process_candidates_from_mongodb(file_content, file_extension):
    # Get MongoDB stats first for debugging
    db_stats = get_mongodb_stats()
    
    # Process job description
    jd_skills = await process_job_description(file_content, file_extension)
    
    if "error" in jd_skills:
        return {"error": jd_skills["error"], "db_stats": db_stats}
    
    # Get candidates from MongoDB
    candidates = get_candidates_from_mongodb()
    
    if isinstance(candidates, dict) and "error" in candidates:
        return {"error": candidates["error"], "db_stats": db_stats}
    
    # Match each candidate
    matched_candidates = []
    updated_count = 0
    
    for candidate in candidates:
        original_id = candidate.get('_id')  # Save original ID for database update
        matched_candidate = match_candidate(jd_skills, candidate)
        
        # Update candidate in MongoDB with match results
        if not isinstance(matched_candidate, dict) or "error" not in matched_candidate:
            update_success = update_candidate_in_mongodb(original_id, matched_candidate)
            if update_success:
                matched_candidate["db_updated"] = True
                updated_count += 1
        
        matched_candidates.append(matched_candidate)
    
    # Sort candidates by match score (descending)
    sorted_candidates = sorted(matched_candidates, key=lambda x: x.get("match", 0), reverse=True)
    
    return {
        "job_description_skills": jd_skills,
        "candidates_count": len(sorted_candidates),
        "candidates_updated": updated_count,
        "db_stats": db_stats,
        "matched_candidates": sorted_candidates
    }

@app.get("/employees")
def get_employees():
    employees = list(employees_collection.find({}))  # Don't exclude _id
    
    # Convert ObjectId to string and add a sequential id for frontend
    for i, employee in enumerate(employees):
        employee["_id"] = str(employee["_id"])  # Convert ObjectId to string
        employee["id"] = i + 1  # Add sequential id for frontend if needed
    
    return employees
@app.get("/employees/{employee_id}")
def get_employee(employee_id: str):
    employee = employees_collection.find_one({"_id": ObjectId(employee_id)})
    if employee:
        employee["_id"] = str(employee["_id"])  # Convert ObjectId to string
        return employee
    return {"error": "Employee not found"}

@app.post("/employees")
def add_employee(employee: dict):
    result = employees_collection.insert_one(employee)
    return {"message": "Employee added", "id": str(result.inserted_id)}

@app.put("/employees/{employee_id}")
def update_employee(employee_id: str, updated_data: dict):
    try:
        # Remove _id field if present in updated_data to avoid the immutable field error
        if '_id' in updated_data:
            del updated_data['_id']
            
        # Try to use the ID as an ObjectId
        result = employees_collection.update_one({"_id": ObjectId(employee_id)}, {"$set": updated_data})
        if result.modified_count:
            return {"message": "Employee updated"}
        return {"error": "No changes made"}
    except bson.errors.InvalidId:
        # If it's not a valid ObjectId, try to find by the numeric id field
        try:
            employee_id_int = int(employee_id)
            employee = employees_collection.find_one({"id": employee_id_int})
            if employee:
                # Remove _id field if present in updated_data
                if '_id' in updated_data:
                    del updated_data['_id']
                    
                result = employees_collection.update_one({"_id": employee["_id"]}, {"$set": updated_data})
                if result.modified_count:
                    return {"message": "Employee updated"}
            return {"error": "Employee not found"}
        except (ValueError, TypeError):
            return {"error": "Invalid employee ID format"}

# Delete an employee
@app.delete("/employees/{employee_id}")
def delete_employee(employee_id: str):
    result = employees_collection.delete_one({"_id": ObjectId(employee_id)})
    if result.deleted_count:
        return {"message": "Employee deleted"}
    return {"error": "Employee not found"}

# Process document and add employee from resume/CV
@app.post("/process-document")
async def process_document(file: UploadFile = File(...)):
    try:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
            # Copy the uploaded file to the temporary file
            shutil.copyfileobj(file.file, temp_file)
            temp_file_path = temp_file.name
        
        # Process the document
        processor = DocumentProcessor()
        extracted_info = processor.process_document(temp_file_path)
        
        # Save to MongoDB
        inserted_id = processor.save_to_mongodb(extracted_info)
        
        # Clean up the temporary file
        os.unlink(temp_file_path)
        
        # Return the extracted information and MongoDB ID
        return {
            "success": True,
            "extracted_info": extracted_info,
            "inserted_id": inserted_id
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e)
        }

@app.get("/")
async def root():
    return {"message": "Job Description Skills Matcher API is running"}

@app.get("/api/db-status")
async def check_db_status():
    return get_mongodb_stats()

@app.get("/api/employees")
async def get_employees():
    candidates = get_candidates_from_mongodb()
    if isinstance(candidates, dict) and "error" in candidates:
        raise HTTPException(status_code=500, detail=candidates["error"])
    return candidates

@app.post("/api/match")
async def match_job_description(file: UploadFile = File(...)):
    try:
        file_content = await file.read()
        file_extension = os.path.splitext(file.filename)[1]
        
        if file_extension.lower() not in ['.pdf', '.docx', '.doc']:
            raise HTTPException(status_code=400, detail="Unsupported file format. Please upload a PDF or Word document.")
            
        results = await batch_process_candidates_from_mongodb(file_content, file_extension)
        
        if "error" in results:
            raise HTTPException(status_code=500, detail=results["error"])
            
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)